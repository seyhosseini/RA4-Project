#Before using this as the standard must change a a lot of variables including epochs

print("Importing Modules")
import os
from monai.transforms import (Spacingd, Compose, LoadImaged, EnsureChannelFirstd, RandSpatialCropd, CenterSpatialCropd,\
                               RandRotate90d, ToTensord, ScaleIntensityRanged, SpatialPadd)
from monai.networks.nets import UNet
from monai.losses import DiceLoss
from monai.data import Dataset, DataLoader, pad_list_data_collate
print("Monai imported. Importing Torch")
import matplotlib.pyplot as plt
import torch
from torch.optim import Adam
from torch.optim.lr_scheduler import CosineAnnealingLR
import time 
import numpy as np
import docx
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import socket
import matplotlib.pyplot as plt
from mpl_toolkits.mplot3d.art3d import Poly3DCollection
from skimage import measure
from scipy.ndimage import zoom
import nrrd

# FUNCTIONS
def create_dataset(data_dir):
    data_dicts = []
    for filename in os.listdir(data_dir):
        if filename.endswith("_Vx3.nrrd"):
            image_path = os.path.join(data_dir, filename)
            label_filename = filename.replace("_Vx3.nrrd", "_Label.nrrd")
            label_path = os.path.join(data_dir, label_filename)
            data_dicts.append({'image': image_path, 'label': label_path})
    return data_dicts

def plot_3d_object(ax, data, value, color, opacity):
    zoom_factor = 0.5
    mask = data == value
    verts, faces, _, _ = measure.marching_cubes(mask, level=0)
    verts *= 1 / zoom_factor  # Adjust for downsampling
    mesh = Poly3DCollection(verts[faces], alpha=opacity)
    mesh.set_facecolor(color)
    ax.add_collection3d(mesh)

def visualize_and_save_train(inputs, outputs, labels, batch_indext, epoch):
    color_map = np.array([
        [255, 255, 255],  # White for background (0)
        [0, 151, 206],      # Blue for arteries (1)
        [216, 101, 79]       # Red for veins (2)
    ], dtype=np.uint8)
    saved_train_image_paths = []

    with open(md_file_path, "a") as md_file:
        for i in range(outputs.shape[0]):  # Iterate over the batch dimension
            if i < inputs.shape[0]:
                fig, ax = plt.subplots(1, 5, figsize=(25, 5))
                # Get the input image and normalize it for visualization
                input_slice = inputs[i, 0, :, :, :].detach().cpu().numpy()
                input_slice_mid = input_slice[input_slice.shape[0] // 2]  # Middle slice of the input volume
                # input_slice_mid_normalized = (input_slice_mid - clippingmin) / \
                #                             (clippingmax - clippingmin)
                zoom_factor = 0.5
                # Get the label and prediction volumes
                label_volume = labels[i, 0].cpu().numpy().astype(np.uint8)
                output_volume = outputs[i].detach().cpu().numpy()
                predicted_labels_volume = np.argmax(output_volume, axis=0).astype(np.uint8)
                downsampled_label = zoom(label_volume, zoom_factor, order=0)
                ax[3] = fig.add_subplot(1, 5, 4, projection='3d')
                plot_3d_object(ax[3], downsampled_label, 1, (0, 0.592157, 0.807843), 1)  # Example values
                plot_3d_object(ax[3], downsampled_label, 2, (0.847059, 0.396078, 0.309804), 1)
                ax[3].set_title("3D Ground Truth")
                
                # Select the middle slice index after argmax reduction
                slice_idx = predicted_labels_volume.shape[0] // 2
                label_slice = label_volume[slice_idx]
                predicted_labels_slice = predicted_labels_volume[slice_idx]
                downsampled_prediction = zoom(predicted_labels_volume, zoom_factor, order=0)
                ax[4] = fig.add_subplot(1, 5, 5, projection='3d')
                plot_3d_object(ax[4], downsampled_prediction, 1, (0, 0.592157, 0.807843), .5)
                plot_3d_object(ax[4], downsampled_prediction, 2, (0.847059, 0.396078, 0.309804), .5)
                ax[4].set_title("3D Prediction")
                
                non_zero_coords = np.argwhere(downsampled_label)
                min_coords = non_zero_coords.min(axis=0) * 1/zoom_factor
                max_coords = non_zero_coords.max(axis=0) * 1/zoom_factor

                ax[3].set_xlim(min_coords[0], max_coords[0])
                ax[3].set_ylim(min_coords[1], max_coords[1])
                ax[3].set_zlim(min_coords[2], max_coords[2])
                ax[4].set_xlim(min_coords[0], max_coords[0])
                ax[4].set_ylim(min_coords[1], max_coords[1])
                ax[4].set_zlim(min_coords[2], max_coords[2])

                # Remove the labels and ticks for the 3D ground truth plot
                ax[3].set_xticklabels([])
                ax[3].set_yticklabels([])
                ax[3].set_zticklabels([])
                # ax[3].set_xticks([])
                # ax[3].set_yticks([])
                # ax[3].set_zticks([])
                # ax[3].axis('off')

                # Remove the labels and ticks for the 3D prediction plot
                ax[4].set_xticklabels([])
                ax[4].set_yticklabels([])
                ax[4].set_zticklabels([])
                # ax[4].set_xticks([])
                # ax[4].set_yticks([])
                # ax[4].set_zticks([])
                # ax[4].axis('off')

                # Apply the color map to the label and prediction slices
                label_rgb_slice = color_map[label_slice]
                prediction_rgb_slice = color_map[predicted_labels_slice]

                # for a in [ax[3], ax[4]]:
                #     a.set_axis_off()
                # Create the plots

                ax[0].imshow(input_slice_mid, cmap='gray')
                ax[0].set_title("Input")
                ax[1].imshow(label_rgb_slice)
                ax[1].set_title("Ground Truth")
                ax[2].imshow(prediction_rgb_slice)
                ax[2].set_title("Prediction")
                
                fig.suptitle(f'Training Image: Epoch-{epoch}-Batch-{batch_indext}-Image-{i}', fontsize=16)
                img_path = f"./PNG/T-Epoch-{epoch}-Batch-{batch_indext}-Image-{i}.png"
                print(f"T-Epoch-{epoch}-Batch-{batch_indext}-Image-{i}.png")

                # Save the image
                plt.savefig(img_path)
                plt.close()
                saved_train_image_paths.append(img_path)
                md_file.write(f"![T-Epoch-{epoch}-Batch-{batch_indext}-Image-{i}.png]({img_path})\n\n")

    print(f"Markdown file updated at {md_file_path}")
    return saved_train_image_paths

def visualize_and_save_valid(inputs, outputs, labels, batch_indexv, epoch):
    color_map = np.array([
        [255, 255, 255],  # White for background (0)
        [0, 151, 206],      # Blue for arteries (1)
        [216, 101, 79]       # Red for veins (2)
    ], dtype=np.uint8)
    saved_valid_image_paths = []

    with open(md_file_path, "a") as md_file:
        for i in range(outputs.shape[0]):  # Iterate over the batch dimension
            if i < inputs.shape[0]:
                fig, ax = plt.subplots(1, 5, figsize=(25, 5))
                zoom_factor = 0.5
                # Get the input image and normalize it for visualization
                input_slice = inputs[i, 0, :, :, :].detach().cpu().numpy()
                input_slice_mid = input_slice[input_slice.shape[0] // 2]  # Middle slice of the input volume
                label_volume = labels[i, 0].cpu().numpy().astype(np.uint8)
                output_volume = outputs[i].detach().cpu().numpy()
                predicted_labels_volume = np.argmax(output_volume, axis=0).astype(np.uint8)
                downsampled_label = zoom(label_volume, zoom_factor, order=0)
                ax[3] = fig.add_subplot(1, 5, 4, projection='3d')
                plot_3d_object(ax[3], downsampled_label, 1, (0, 0.592157, 0.807843), 1)  # Example values
                plot_3d_object(ax[3], downsampled_label, 2, (0.847059, 0.396078, 0.309804), 1)
                ax[3].set_title("3D Ground Truth")

                # Select the middle slice index after argmax reduction
                slice_idx = predicted_labels_volume.shape[0] // 2
                label_slice = label_volume[slice_idx]
                predicted_labels_slice = predicted_labels_volume[slice_idx]

                # Apply the color map to the label and prediction slices
                label_rgb_slice = color_map[label_slice]
                prediction_rgb_slice = color_map[predicted_labels_slice]
                downsampled_prediction = zoom(predicted_labels_volume, zoom_factor, order=0)
                ax[4] = fig.add_subplot(1, 5, 5, projection='3d')
                plot_3d_object(ax[4], downsampled_prediction, 1, (0, 0.592157, 0.807843), .5)
                plot_3d_object(ax[4], downsampled_prediction, 2, (0.847059, 0.396078, 0.309804), .5)
                ax[4].set_title("3D Prediction")
                non_zero_coords = np.argwhere(downsampled_label)
                min_coords = non_zero_coords.min(axis=0) * 1/zoom_factor
                max_coords = non_zero_coords.max(axis=0) * 1/zoom_factor

                ax[3].set_xlim(min_coords[0], max_coords[0])
                ax[3].set_ylim(min_coords[1], max_coords[1])
                ax[3].set_zlim(min_coords[2], max_coords[2])
                ax[4].set_xlim(min_coords[0], max_coords[0])
                ax[4].set_ylim(min_coords[1], max_coords[1])
                ax[4].set_zlim(min_coords[2], max_coords[2])

                # Remove the labels and ticks for the 3D ground truth plot
                ax[3].set_xticklabels([])
                ax[3].set_yticklabels([])
                ax[3].set_zticklabels([])

                # ax[3].set_xticks([])
                # ax[3].set_yticks([])
                # ax[3].set_zticks([])
                # ax[3].axis('off')

                # Remove the labels and ticks for the 3D prediction plot
                ax[4].set_xticklabels([])
                ax[4].set_yticklabels([])
                ax[4].set_zticklabels([])

                # ax[4].set_xticks([])
                # ax[4].set_yticks([])
                # ax[4].set_zticks([])
                # ax[4].axis('off')
                # for a in [ax[3], ax[4]]:
                #     a.set_axis_off()
                # Create the plots

                ax[0].imshow(input_slice_mid, cmap='gray')
                ax[0].set_title("Input")
                ax[1].imshow(label_rgb_slice)
                ax[1].set_title("Ground Truth")
                ax[2].imshow(prediction_rgb_slice)
                ax[2].set_title("Prediction")
                
                fig.suptitle(f'Validation Image: Epoch-{epoch}-Batch-{batch_indexv}-Image-{i}', fontsize=16)
                img_path = f"./PNG/V-Epoch-{epoch}-Batch-{batch_indexv}-Image-{i}.png"
                print(f"V-Epoch-{epoch}-Batch-{batch_indexv}-Image-{i}.png")
                # Save the image
                plt.savefig(img_path)
                plt.close()
                saved_valid_image_paths.append(img_path)
                # Write to the markdown file
                md_file.write(f"![V-Epoch-{epoch}-Batch-{batch_indexv}-Image-{i}.png]({img_path})\n\n")

    print(f"Markdown file updated at {md_file_path}")

    return saved_valid_image_paths

def select_images(image_paths, num_select=5):
    selected_indices = np.linspace(0, len(image_paths) - 1, num_select, dtype=int)
    return [image_paths[i] for i in selected_indices]

def plot_losses(train_losses, val_losses):
    plt.figure(figsize=(10, 5))
    plt.plot(train_losses, label='Train Loss')
    plt.plot(val_losses, label='Validation Loss')
    plt.xlabel('Epochs')
    plt.ylabel('Loss')
    plt.title('Training and Validation Losses')
    plt.legend()
    plt.grid(True, which='both', linestyle='-', linewidth=0.5, color="black")
    plt.minorticks_on()
    plt.grid(True, which='minor', linestyle='-', linewidth=0.3, color='gray')
    plt.draw()
    plt.pause(0.001)
    plt.savefig("./PNG/lossfunction.png")
    plt.close()
# Loss Plotting Function
def plot_losses_log(train_losses, val_losses):
    plt.figure(figsize=(10, 5))
    plt.plot(train_losses, label='Train Loss')
    plt.plot(val_losses, label='Validation Loss')
    plt.xlabel('Epochs')
    plt.ylabel('Loss')
    plt.yscale('log')
    plt.title('Training and Validation Losses')
    plt.legend()
    plt.grid(True, which='both', linestyle='-', linewidth=0.5, color="black")
    plt.minorticks_on()
    plt.grid(True, which='minor', linestyle='-', linewidth=0.3, color='gray')
    plt.draw()
    plt.pause(0.001)
    plt.savefig("./PNG/log-lossfunction.png")
    plt.close()

def plot_learning_rate(learning_rate):
    epochs = range(1, len(train_losses) + 1) 
    plt.figure(figsize=(10, 5))
    plt.plot(epochs, learning_rate, label='Learning Rate')
    plt.xlabel('Epochs')
    plt.ylabel('Learning Rate')
    plt.title('Learning Rate Change Vs Epoch')
    plt.legend()
    plt.grid(True, which='both', linestyle='-', linewidth=0.5, color="black")
    plt.minorticks_on()
    plt.grid(True, which='minor', linestyle='-', linewidth=0.3, color='gray')
    plt.draw()
    plt.pause(0.001)
    plt.savefig("./PNG/learning_rate.png")
    plt.close()

def plot_log_learning_rate(learning_rate):
    epochs = range(1, len(learning_rate) + 1) 
    plt.figure(figsize=(10, 5))
    plt.plot(epochs, learning_rate, label='Learning Rate')
    plt.xlabel('Epochs')
    plt.ylabel('Learning Rate')
    plt.yscale('log')
    plt.title('Learning Rate Change Vs Epoch')
    plt.legend()
    plt.grid(True, which='both', linestyle='-', linewidth=0.5, color="black")
    plt.minorticks_on()
    plt.grid(True, which='minor', linestyle='-', linewidth=0.3, color='gray')
    plt.draw()
    plt.pause(0.001)
    plt.savefig("./PNG/log_learning_rate.png")
    plt.close()

def train_step(batch_data, model, loss_function, optimizer, device):
    # print("Reading in the images")
    # print(f"Processing filenames: {batch_data['filename']}")
    images, labels = batch_data['image'], batch_data['label']
    images, labels = images.to(device), labels.to(device)
    optimizer.zero_grad()
    print("Passing Through the Network")
    outputs = model(images)

    print(f"Output shape: {outputs.shape}, Label shape: {labels.shape}")
    print("Label tensor values (sample):", labels[0, :, 0, 0, 0])
    # if torch.any(labels == 3):
        # print(f"Filename with label 3: {batch_data['filename']}")

    loss = loss_function(outputs, labels)
    print("Loss Backward")
    loss.backward()
    print("Stepping Optimizer")
    optimizer.step()
    return loss.item(), outputs

def create_word_document(expdesc, log_loss_path, selected_train_images, selected_val_images, learning_rate_path, log_learning_rate_path, final_time):
    ########################################################################
    parent_directory = os.path.dirname(os.getcwd())
    output_path = os.path.join(parent_directory, 'JobReport.docx')

    # Check if the file exists. If yes, open the existing file; if no, create a new one
    if os.path.exists(output_path):
        doc = docx.Document(output_path)
        print(f"Trying to append to existing document: {output_path}")
    else:
        doc = docx.Document()
        print(f"Creating new document: {output_path}")
    print("Appending Images Now")
    # Append content to the document
    doc.add_paragraph("Experiment Description:")
    doc.add_paragraph(expdesc)
    doc.add_paragraph(final_time)
    doc.add_picture(log_loss_path, width=Inches(6))
    doc.add_picture(learning_rate_path, width=Inches(6))
    doc.add_picture(log_learning_rate_path, width=Inches(6))
    for img_path in selected_train_images + selected_val_images:
        doc.add_picture(img_path, width=Inches(6))


    # Save the document
    doc.save(output_path)

    # Print the path where the file was updated or created
    print(f"Document updated/created: {output_path}")
    ############################################################################

# All Constants
md_file_path = "./visualization.md"
hostname = socket.gethostname()
expdesc = (f"Job Test: \n Copy of Job 8 \n Train Images: 2 \n Validation Images: 5 \n Train Crop: RandSpatialCrop \n \
Validation Crop: RandSpatialCrop \n Epochs: 1500 \n Learning Rate: 1e-3 \n \
Spacing: (0.8, 0.8, 0.8) \n Batch Size: 4 \n Patch Size: (128, 128, 128) \n Compute Node: {hostname}") # REMEBER TO CHANGE THIS EACH JOB

with open('expdesc.txt', 'w') as file:
    file.write(expdesc)


log_loss_path = "./PNG/log-lossfunction.png"
learning_rate_path = "./PNG/learning_rate.png"
log_learning_rate_path ="./PNG/log_learning_rate.png"
train_data_dir = "../../TrainData/"
val_data_dir = "../../ValidationData/"
model_save_path = "./model/Nate_Unet.pth"
optimizer_save_path = "./model/Nate_Unet_optimizer.pth"
train_image_paths = []
val_image_paths = []
train_losses = []
val_losses = []

word_output_path = "./JobReport.docx"
# Main training loop
num_epochs = 50
display_interval = 1 
starttime = time.time()



print("Defining the create_dataset function...")
# Function to create dataset
print("Setting data paths...")
# Set data paths

print("Creating datasets...")
train_files = create_dataset(train_data_dir)
val_files = create_dataset(val_data_dir)

downsampling_transform = Spacingd(
    keys=['image', 'label'], 
    pixdim=(0.8, 0.8, 0.8), 
    mode=('bilinear', 'nearest') ###############################################################
)

# Define the size of the cropped region
roi_size = (128, 128, 128)
# roi_size = (64, 64, 64)
print("Defining transformations...")
# Transformations
print(f"Elapsed: {(time.time()-starttime)/60:.2f} Minutes")
train_transforms = Compose([
    LoadImaged(keys=['image', 'label']),
    EnsureChannelFirstd(keys=['image', 'label']),
    downsampling_transform, #####################################################################
    # SpatialPadd(keys=['image', 'label'], spatial_size=roi_size, method='symmetric'),  # Add padding here
    RandSpatialCropd(keys=['image', 'label'], roi_size=roi_size),
    RandRotate90d(keys=['image', 'label'], prob=0.5),
    ScaleIntensityRanged(keys=['image'], a_min=-1100, a_max=3000, b_min=-1.0, b_max=1.0, clip=True),
    ToTensord(keys=['image', 'label']),
])

val_transforms = Compose([
    LoadImaged(keys=['image', 'label']),
    EnsureChannelFirstd(keys=['image', 'label']),
    downsampling_transform,
    # SpatialPadd(keys=['image', 'label'], spatial_size=roi_size, method='symmetric'),  # Add padding here
    RandSpatialCropd(keys=['image', 'label'], roi_size=roi_size),
    ScaleIntensityRanged(keys=['image'], a_min=-1100, a_max=3000, b_min=-1.0, b_max=1.0, clip=True),
    ToTensord(keys=['image', 'label']),
])

print("Initializing data loaders...")
# Data Loaders
train_ds = Dataset(train_files, train_transforms)
val_ds = Dataset(val_files, val_transforms)
train_loader = DataLoader(train_ds, batch_size=4, shuffle=True) #collate_fn=pad_list_data_collate)
val_loader = DataLoader(val_ds, batch_size=4, shuffle=True) #collate_fn=pad_list_data_collate)

print("Initializing U-Net model...")
# Model
net = UNet(
    spatial_dims=3,
    in_channels=1,
    out_channels=3,
    channels=(16, 32, 64, 128, 256),
    strides=(2, 2, 2, 2),
    num_res_units=2,
)

# if torch.cuda.is_available():
#     device = torch.device("cuda")
#     num_gpus = torch.cuda.device_count()
#     print(f"Number of available GPUs: {num_gpus}") # USE torch.nn.DistributedDataParallel(net)
#     if num_gpus > 1:
#         net = torch.nn.DataParallel(net)  # Wrap the model with DataParallel
#     net.to(device)
# else:
#     device = torch.device("cpu")
#     net.to(device)

print(f"Elapsed: {(time.time()-starttime)/60:.2f} Minutes")
loss_function = DiceLoss(to_onehot_y=True, softmax=True) #########################################
criterion = loss_function
optimizer = Adam(net.parameters(), 1e-3)
scheduler = CosineAnnealingLR(optimizer, T_max=1500, eta_min=1e-6)
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
net.to(device)

if not os.path.exists("./model"):
    os.makedirs("./model", exist_ok=True)
if os.path.exists(model_save_path):
    net.load_state_dict(torch.load(model_save_path, map_location=device))
    print("Model state loaded successfully.")
if os.path.exists(optimizer_save_path):
    optimizer.load_state_dict(torch.load(optimizer_save_path))
    print("Optimizer state loaded successfully.")

print("Setting up loss function and optimizer...")
learning_rates = []
print("Starting the training loop...")
for epoch in range(num_epochs):
    epoch_train_loss = 0.0
    epoch_val_loss = 0.0
    
    # Training Phase
    net.train()
    # batch_indext = 0
    print(f"Elapsed: {(time.time()-starttime)/60:.2f} Minutes")
    print("\nReading in the first batch - T ..")
    for batch_indext, batch_data in enumerate(train_loader):
        print(f"Training Iteration: {batch_indext}")
        loss, output = train_step(batch_data, net, criterion, optimizer, device)
        epoch_train_loss += loss
        # if iteration % display_interval == 0:
        if (epoch + 1) % 1 == 0 or epoch == 0:
            train_image_batch_paths = visualize_and_save_train(batch_data['image'], output, batch_data['label'], batch_indext, epoch) 
        # train_image_batch_paths = visualize_and_save_train(batch_data['image'], output, batch_data['label'], epoch, iteration)
        train_image_paths.extend(train_image_batch_paths)

        print("\nReading in the next batch - T [if any] ..")
        # batch_indext = batch_indext + 1
    train_losses.append(epoch_train_loss / len(train_loader))
    print("Running the validation loop ..")
    
    # Validation Phase
    net.eval()
    # batch_indexv = 0
    with torch.no_grad():
        print("Reading in the first batch - V ..")
        for batch_indexv, batch_data in enumerate(val_loader):
            print(f"Validation Iteration: {batch_indexv}")
            # print(f"Processing filenames: {batch_data['filename']}")
            images, labels = batch_data['image'], batch_data['label']
            images, labels = images.to(device), labels.to(device)
            outputs = net(images)

            print(f"Validation - Output shape: {outputs.shape}, Label shape: {labels.shape}")

            loss = criterion(outputs, labels)
            epoch_val_loss += loss.item()
            # if iteration % display_interval == 0:
            if (epoch + 1) % 1 == 0 or epoch == 0:
                val_image_batch_paths = visualize_and_save_valid(batch_data['image'], outputs, batch_data['label'], batch_indexv, epoch) 
            # val_image_batch_paths = visualize_and_save_valid(batch_data['image'], outputs, batch_data['label'], epoch, iteration)
                val_image_paths.extend(val_image_batch_paths)
            print("Reading in the next batch - V [if any] ..")
            # batch_indexv = batch_indexv + 1
        print()    
    
    val_losses.append(epoch_val_loss / len(val_loader))
    scheduler.step()
    for param_group in optimizer.param_groups:
        current_lr = param_group['lr']
        print(f"Epoch: {epoch}, Current Learning Rate: {current_lr}")
    current_lr = optimizer.param_groups[0]['lr']
    learning_rates.append(current_lr)
    # Plot Losses
    print("Plotting losses ..")
    plot_losses(train_losses, val_losses)
    plot_losses_log(train_losses, val_losses)
    print(f"Epoch [{epoch+1}/{num_epochs}] - Train Loss: {train_losses[-1]:.4f}, Val Loss: {val_losses[-1]:.4f}")
    plot_learning_rate(learning_rates)
    plot_log_learning_rate(learning_rates)

selected_train_images = select_images(train_image_paths)
selected_val_images = select_images(val_image_paths)

# Call the function to create the document

torch.save(net.state_dict(), model_save_path)
torch.save(optimizer.state_dict(), optimizer_save_path)
print(f"Elapsed: {(time.time()-starttime)/60:.2f} Minutes")
elapsed_minutes = (time.time() - starttime) / 60
final_time = f"Time: {elapsed_minutes:.2f} Minutes"
create_word_document(expdesc, log_loss_path, selected_train_images, selected_val_images, learning_rate_path, log_learning_rate_path, final_time)